"""
Script de génération du document de présentation NFC Pro
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Définit la couleur de fond d'une cellule"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(paragraph):
    """Ajoute une ligne horizontale"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2563EB')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_document():
    doc = Document()

    # Configuration des marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ==========================================
    # PAGE DE GARDE
    # ==========================================

    # Espacement haut
    for _ in range(4):
        doc.add_paragraph()

    # Logo/Titre principal
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("📱 NFC PRO")
    run.bold = True
    run.font.size = Pt(48)
    run.font.color.rgb = RGBColor(37, 99, 235)  # Bleu primaire

    # Sous-titre
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Application Mobile Professionnelle\nNFC & RFID")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(75, 85, 99)

    doc.add_paragraph()

    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Lisez • Écrivez • Copiez • Émulez • Partagez")
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = RGBColor(107, 114, 128)

    # Espacement
    for _ in range(6):
        doc.add_paragraph()

    # Type de document
    doc_type = doc.add_paragraph()
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_type.add_run("DOCUMENT DE PRÉSENTATION")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(37, 99, 235)

    # Version et date
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Version 1.0 — Novembre 2025")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(107, 114, 128)

    # Confidentialité
    for _ in range(4):
        doc.add_paragraph()

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run("CONFIDENTIEL")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(239, 68, 68)

    # Saut de page
    doc.add_page_break()

    # ==========================================
    # SOMMAIRE
    # ==========================================

    toc_title = doc.add_paragraph()
    run = toc_title.add_run("SOMMAIRE")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(37, 99, 235)
    add_horizontal_line(toc_title)

    doc.add_paragraph()

    toc_items = [
        ("1.", "Résumé Exécutif", "3"),
        ("2.", "Présentation du Projet", "4"),
        ("3.", "Fonctionnalités Clés", "6"),
        ("4.", "Architecture Technique", "9"),
        ("5.", "Modèle Économique", "11"),
        ("6.", "Planning & Investissement", "13"),
        ("7.", "Pourquoi NFC Pro ?", "15"),
    ]

    for num, title_text, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)

        run_num = p.add_run(num + " ")
        run_num.bold = True
        run_num.font.size = Pt(12)
        run_num.font.color.rgb = RGBColor(37, 99, 235)

        run_title = p.add_run(title_text)
        run_title.font.size = Pt(12)

        run_dots = p.add_run(" " + "." * 50 + " ")
        run_dots.font.size = Pt(12)
        run_dots.font.color.rgb = RGBColor(209, 213, 219)

        run_page = p.add_run(page)
        run_page.font.size = Pt(12)
        run_page.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_page_break()

    # ==========================================
    # 1. RÉSUMÉ EXÉCUTIF
    # ==========================================

    section_title = doc.add_paragraph()
    run = section_title.add_run("1. RÉSUMÉ EXÉCUTIF")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(37, 99, 235)
    add_horizontal_line(section_title)

    doc.add_paragraph()

    # Introduction
    intro = doc.add_paragraph()
    run = intro.add_run("NFC Pro")
    run.bold = True
    intro.add_run(" est une application mobile professionnelle multiplateforme (Android & iOS) conçue pour répondre aux besoins des professionnels en matière de technologie NFC et RFID.")

    doc.add_paragraph()

    # Points clés en tableau
    key_points_title = doc.add_paragraph()
    run = key_points_title.add_run("Points clés du projet")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    data = [
        ("🎯 Marché cible", "Professionnels IT, entreprises, développeurs IoT"),
        ("💰 Modèle", "Freemium + Abonnement Pro à 19€/an"),
        ("📱 Plateformes", "Android & iOS (Flutter)"),
        ("☁️ Infrastructure", "Microsoft Azure (serverless)"),
        ("🔐 Sécurité", "Chiffrement AES-256, conformité RGPD"),
    ]

    for i, (key, value) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        set_cell_shading(row.cells[0], 'F3F4F6')
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Proposition de valeur
    value_title = doc.add_paragraph()
    run = value_title.add_run("Proposition de valeur unique")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(17, 24, 39)

    values = [
        "Support exhaustif de toutes les puces NFC/RFID du marché",
        "Solution tout-en-un : lecture, écriture, copie, émulation, cartes de visite",
        "Intégration native Google Wallet et Apple Wallet",
        "Synchronisation cloud multi-appareils",
        "Interface professionnelle optimisée pour l'efficacité",
    ]

    for v in values:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(v)

    doc.add_paragraph()

    # Objectifs financiers
    fin_title = doc.add_paragraph()
    run = fin_title.add_run("Objectifs à 12 mois")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Table Grid'

    objectives = [
        ("Téléchargements", "30 000"),
        ("Abonnés Pro", "900"),
        ("Revenus annuels", "17 100 €"),
        ("Note moyenne stores", "4.5 / 5"),
    ]

    for i, (metric, target) in enumerate(objectives):
        row = table2.rows[i]
        row.cells[0].text = metric
        row.cells[1].text = target
        set_cell_shading(row.cells[0], 'DBEAFE')
        row.cells[1].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ==========================================
    # 2. PRÉSENTATION DU PROJET
    # ==========================================

    section_title = doc.add_paragraph()
    run = section_title.add_run("2. PRÉSENTATION DU PROJET")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(37, 99, 235)
    add_horizontal_line(section_title)

    doc.add_paragraph()

    # Contexte
    ctx_title = doc.add_paragraph()
    run = ctx_title.add_run("Contexte et opportunité")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    ctx = doc.add_paragraph()
    ctx.add_run("La technologie NFC (Near Field Communication) connaît une croissance exponentielle avec l'adoption massive des paiements sans contact, des badges d'accès intelligents et de l'IoT. Le marché des applications NFC professionnelles reste cependant fragmenté, avec des solutions souvent incomplètes ou peu intuitives.")

    doc.add_paragraph()

    ctx2 = doc.add_paragraph()
    run = ctx2.add_run("NFC Pro")
    run.bold = True
    ctx2.add_run(" se positionne comme LA solution tout-en-un pour les professionnels, combinant puissance fonctionnelle et simplicité d'utilisation.")

    doc.add_paragraph()

    # Public cible
    target_title = doc.add_paragraph()
    run = target_title.add_run("Public cible")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table3 = doc.add_table(rows=5, cols=3)
    table3.style = 'Table Grid'

    # En-tête
    header_row = table3.rows[0]
    headers = ["Segment", "Profil", "Besoins"]
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        set_cell_shading(header_row.cells[i], '2563EB')
        for run in header_row.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    targets = [
        ("🔒 Sécurité IT", "RSSI, pentesters, auditeurs", "Analyse tags, tests de sécurité"),
        ("🏢 Entreprises", "Responsables logistique, RH", "Inventaire, contrôle d'accès"),
        ("💻 Développeurs", "Dev IoT, intégrateurs", "Prototypage, tests, debug"),
        ("🤝 Commercial", "Commerciaux, marketeurs", "Cartes de visite digitales"),
    ]

    for i, (segment, profil, besoins) in enumerate(targets):
        row = table3.rows[i + 1]
        row.cells[0].text = segment
        row.cells[1].text = profil
        row.cells[2].text = besoins

    doc.add_paragraph()
    doc.add_paragraph()

    # Analyse concurrentielle
    comp_title = doc.add_paragraph()
    run = comp_title.add_run("Positionnement concurrentiel")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table4 = doc.add_table(rows=5, cols=5)
    table4.style = 'Table Grid'

    # En-tête
    headers2 = ["Fonctionnalité", "NFC Pro", "NFC Tools", "TagWriter", "TagInfo"]
    header_row2 = table4.rows[0]
    for i, h in enumerate(headers2):
        header_row2.cells[i].text = h
        set_cell_shading(header_row2.cells[i], '2563EB')
        for run in header_row2.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    comp_data = [
        ("Lecture/Écriture", "✅ Complet", "✅", "✅", "✅"),
        ("Émulation HCE", "✅", "⚠️ Limité", "❌", "❌"),
        ("Cartes de visite", "✅ Intégré", "❌", "❌", "❌"),
        ("Google/Apple Wallet", "✅", "❌", "❌", "❌"),
    ]

    for i, row_data in enumerate(comp_data):
        row = table4.rows[i + 1]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
            if j == 1:  # Colonne NFC Pro
                set_cell_shading(row.cells[j], 'DCFCE7')

    doc.add_page_break()

    # ==========================================
    # 3. FONCTIONNALITÉS CLÉS
    # ==========================================

    section_title = doc.add_paragraph()
    run = section_title.add_run("3. FONCTIONNALITÉS CLÉS")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(37, 99, 235)
    add_horizontal_line(section_title)

    doc.add_paragraph()

    # Module Lecture
    mod1_title = doc.add_paragraph()
    run = mod1_title.add_run("📖 Module Lecture NFC/RFID")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(16, 185, 129)

    read_features = [
        "Détection automatique du type de tag en moins de 500ms",
        "Lecture complète des données NDEF (URL, texte, vCard, WiFi...)",
        "Affichage du dump mémoire hexadécimal pour analyse avancée",
        "Identification de l'UID et des caractéristiques techniques",
        "Export des données en JSON, XML ou format binaire",
        "Historique des lectures avec synchronisation cloud",
    ]

    for f in read_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f)

    doc.add_paragraph()

    # Module Écriture
    mod2_title = doc.add_paragraph()
    run = mod2_title.add_run("✏️ Module Écriture NFC")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(139, 92, 246)

    write_features = [
        "Écriture de tous les formats NDEF standards",
        "Templates personnalisables et réutilisables",
        "Configuration WiFi et appairage Bluetooth en un tap",
        "Protection par mot de passe et verrouillage permanent",
        "Écriture en lot pour production de masse",
        "Mode avancé pour écriture de données brutes",
    ]

    for f in write_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f)

    doc.add_paragraph()

    # Module Copie
    mod3_title = doc.add_paragraph()
    run = mod3_title.add_run("📋 Module Copie & Backup")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(245, 158, 11)

    copy_features = [
        "Copie complète du contenu NDEF vers un nouveau tag",
        "Backup intégral de la mémoire du tag",
        "Restauration depuis fichier de sauvegarde",
        "Détection automatique de la compatibilité de copie",
        "Avertissements légaux et traçabilité des opérations",
    ]

    for f in copy_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f)

    doc.add_page_break()

    # Module Émulation
    mod4_title = doc.add_paragraph()
    run = mod4_title.add_run("📡 Module Émulation HCE (Android)")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(6, 182, 212)

    emu_features = [
        "Émulation de carte de visite NFC sans tag physique",
        "Profils d'émulation multiples interchangeables",
        "UID configurable pour tests avancés",
        "Widget d'activation rapide",
        "Compatible avec tous les lecteurs NFC standards",
    ]

    for f in emu_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f)

    note = doc.add_paragraph()
    run = note.add_run("Note : ")
    run.bold = True
    run.font.color.rgb = RGBColor(239, 68, 68)
    note.add_run("L'émulation HCE n'est pas disponible sur iOS en raison des restrictions Apple.")

    doc.add_paragraph()

    # Module Cartes de visite
    mod5_title = doc.add_paragraph()
    run = mod5_title.add_run("🪪 Module Cartes de Visite Numériques")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(236, 72, 153)

    card_features = [
        "Création de cartes de visite professionnelles personnalisées",
        "Bibliothèque de templates élégants et modernes",
        "Upload photo de profil et logo entreprise",
        "Intégration des réseaux sociaux (LinkedIn, Twitter, etc.)",
        "Partage multi-canal : NFC, QR Code, lien URL, vCard",
        "Page web publique personnalisée pour chaque carte",
        "Scan OCR des cartes de visite papier",
        "Analytics : nombre de vues, scans, interactions",
    ]

    for f in card_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f)

    doc.add_paragraph()

    # Module Wallet
    mod6_title = doc.add_paragraph()
    run = mod6_title.add_run("💳 Intégration Wallet")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(37, 99, 235)

    wallet_features = [
        "Export vers Google Wallet (Android)",
        "Export vers Apple Wallet (iOS)",
        "Mise à jour dynamique des informations",
        "Notifications push contextuelles",
        "Accès rapide depuis l'écran de verrouillage",
    ]

    for f in wallet_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f)

    doc.add_page_break()

    # Puces supportées
    chips_title = doc.add_paragraph()
    run = chips_title.add_run("🔧 Puces NFC/RFID Supportées")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table5 = doc.add_table(rows=6, cols=3)
    table5.style = 'Table Grid'

    header_row5 = table5.rows[0]
    for i, h in enumerate(["Famille", "Modèles", "Capacité"]):
        header_row5.cells[i].text = h
        set_cell_shading(header_row5.cells[i], '2563EB')
        for run in header_row5.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    chips = [
        ("NTAG (NXP)", "210, 213, 215, 216, 413/424 DNA", "48 - 888 bytes"),
        ("MIFARE Classic", "1K, 4K, EV1", "1 - 4 KB"),
        ("MIFARE DESFire", "EV1, EV2, EV3", "2 - 8 KB"),
        ("MIFARE Ultralight", "Standard, C, EV1", "48 - 192 bytes"),
        ("Autres", "FeliCa, ICODE, ST25", "Variable"),
    ]

    for i, (famille, modeles, capacite) in enumerate(chips):
        row = table5.rows[i + 1]
        row.cells[0].text = famille
        row.cells[1].text = modeles
        row.cells[2].text = capacite
        set_cell_shading(row.cells[0], 'F3F4F6')
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    support_note = doc.add_paragraph()
    run = support_note.add_run("+ de 30 types de puces supportées")
    run.bold = True
    support_note.add_run(", incluant les dernières générations avec authentification AES-128 et Secure Unique NFC (SUN).")

    doc.add_page_break()

    # ==========================================
    # 4. ARCHITECTURE TECHNIQUE
    # ==========================================

    section_title = doc.add_paragraph()
    run = section_title.add_run("4. ARCHITECTURE TECHNIQUE")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(37, 99, 235)
    add_horizontal_line(section_title)

    doc.add_paragraph()

    # Stack Mobile
    stack_title = doc.add_paragraph()
    run = stack_title.add_run("Stack Technique Mobile")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table6 = doc.add_table(rows=7, cols=3)
    table6.style = 'Table Grid'

    header_row6 = table6.rows[0]
    for i, h in enumerate(["Composant", "Technologie", "Justification"]):
        header_row6.cells[i].text = h
        set_cell_shading(header_row6.cells[i], '2563EB')
        for run in header_row6.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    stack_data = [
        ("Framework", "Flutter 3.24+", "Cross-platform, performances natives"),
        ("Langage", "Dart 3.5+", "Typage fort, async natif"),
        ("State Mgmt", "Riverpod 2.5+", "Scalable, testable"),
        ("NFC (Android)", "Kotlin + NFC API", "Accès complet hardware"),
        ("NFC (iOS)", "Swift + CoreNFC", "Respect guidelines Apple"),
        ("Local DB", "Hive + SQLite", "Rapide + relationnel"),
    ]

    for i, (comp, tech, just) in enumerate(stack_data):
        row = table6.rows[i + 1]
        row.cells[0].text = comp
        row.cells[1].text = tech
        row.cells[2].text = just
        set_cell_shading(row.cells[0], 'F3F4F6')

    doc.add_paragraph()

    # Infrastructure Azure
    azure_title = doc.add_paragraph()
    run = azure_title.add_run("Infrastructure Azure")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table7 = doc.add_table(rows=8, cols=3)
    table7.style = 'Table Grid'

    header_row7 = table7.rows[0]
    for i, h in enumerate(["Service", "Usage", "Tier"]):
        header_row7.cells[i].text = h
        set_cell_shading(header_row7.cells[i], '0078D4')  # Azure blue
        for run in header_row7.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    azure_data = [
        ("Azure Functions", "API Backend serverless", "Premium P1v2"),
        ("Cosmos DB", "Base de données NoSQL", "Serverless"),
        ("Blob Storage", "Stockage fichiers/médias", "Standard GRS"),
        ("Azure B2C", "Authentification", "Standard"),
        ("API Management", "Gateway & rate limiting", "Developer"),
        ("SignalR Service", "Temps réel (sync)", "Standard"),
        ("Notification Hubs", "Push notifications", "Standard"),
    ]

    for i, (service, usage, tier) in enumerate(azure_data):
        row = table7.rows[i + 1]
        row.cells[0].text = service
        row.cells[1].text = usage
        row.cells[2].text = tier
        set_cell_shading(row.cells[0], 'E6F2FF')

    doc.add_paragraph()

    # Schéma simplifié
    archi_desc = doc.add_paragraph()
    run = archi_desc.add_run("Architecture simplifiée")
    run.bold = True
    archi_desc.paragraph_format.space_after = Pt(12)

    archi = doc.add_paragraph()
    archi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    archi_text = """
┌─────────────────────────────────────────────────────────────┐
│                    CLIENTS MOBILES                           │
│         ┌───────────────┐    ┌───────────────┐              │
│         │  Android App  │    │    iOS App    │              │
│         │   (Flutter)   │    │   (Flutter)   │              │
│         └───────────────┘    └───────────────┘              │
└─────────────────────────────────────────────────────────────┘
                              │
                              ▼
┌─────────────────────────────────────────────────────────────┐
│                      AZURE CLOUD                             │
│  ┌─────────┐  ┌──────────┐  ┌─────────┐  ┌──────────┐      │
│  │Functions│  │ Cosmos DB│  │  Blob   │  │  B2C     │      │
│  │  (API)  │  │  (Data)  │  │ Storage │  │  (Auth)  │      │
│  └─────────┘  └──────────┘  └─────────┘  └──────────┘      │
└─────────────────────────────────────────────────────────────┘
"""
    run = archi.add_run(archi_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    doc.add_page_break()

    # Sécurité
    sec_title = doc.add_paragraph()
    run = sec_title.add_run("🔐 Sécurité & Conformité")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table8 = doc.add_table(rows=7, cols=2)
    table8.style = 'Table Grid'

    header_row8 = table8.rows[0]
    header_row8.cells[0].text = "Mesure"
    header_row8.cells[1].text = "Implémentation"
    for cell in header_row8.cells:
        set_cell_shading(cell, 'DC2626')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    security_data = [
        ("Chiffrement transit", "TLS 1.3 obligatoire"),
        ("Chiffrement repos", "AES-256 (Azure managed keys)"),
        ("Authentification", "Azure B2C + JWT RS256"),
        ("Stockage sensible", "Keychain (iOS) / EncryptedSharedPrefs (Android)"),
        ("Protection API", "Rate limiting, WAF, certificate pinning"),
        ("Conformité", "RGPD (droits utilisateurs, DPA, registre)"),
    ]

    for i, (mesure, impl) in enumerate(security_data):
        row = table8.rows[i + 1]
        row.cells[0].text = mesure
        row.cells[1].text = impl
        set_cell_shading(row.cells[0], 'FEE2E2')

    doc.add_page_break()

    # ==========================================
    # 5. MODÈLE ÉCONOMIQUE
    # ==========================================

    section_title = doc.add_paragraph()
    run = section_title.add_run("5. MODÈLE ÉCONOMIQUE")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(37, 99, 235)
    add_horizontal_line(section_title)

    doc.add_paragraph()

    # Offres
    offers_title = doc.add_paragraph()
    run = offers_title.add_run("Structure de l'offre")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table9 = doc.add_table(rows=10, cols=3)
    table9.style = 'Table Grid'

    header_row9 = table9.rows[0]
    headers9 = ["Fonctionnalité", "Gratuit", "Pro (19€/an)"]
    for i, h in enumerate(headers9):
        header_row9.cells[i].text = h
        color = '2563EB' if i == 0 else ('6B7280' if i == 1 else '10B981')
        set_cell_shading(header_row9.cells[i], color)
        for run in header_row9.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    offers_data = [
        ("Lecture NFC", "Illimitée", "Illimitée + avancée"),
        ("Écriture", "5/mois", "Illimitée"),
        ("Copie/Clone", "❌", "✅ Complète"),
        ("Émulation HCE", "❌", "✅"),
        ("Historique", "10 tags", "Illimité + cloud"),
        ("Cartes de visite", "1", "Illimitées"),
        ("Wallet export", "❌", "✅"),
        ("Publicités", "Oui", "Non"),
        ("Support", "Communauté", "Email prioritaire"),
    ]

    for i, (feature, free, pro) in enumerate(offers_data):
        row = table9.rows[i + 1]
        row.cells[0].text = feature
        row.cells[1].text = free
        row.cells[2].text = pro
        if "✅" in pro or pro == "Illimitée + avancée":
            set_cell_shading(row.cells[2], 'DCFCE7')

    doc.add_paragraph()

    # Projections
    proj_title = doc.add_paragraph()
    run = proj_title.add_run("Projections financières sur 24 mois")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table10 = doc.add_table(rows=6, cols=5)
    table10.style = 'Table Grid'

    header_row10 = table10.rows[0]
    headers10 = ["Période", "Downloads", "Abonnés Pro", "MRR", "ARR"]
    for i, h in enumerate(headers10):
        header_row10.cells[i].text = h
        set_cell_shading(header_row10.cells[i], '2563EB')
        for run in header_row10.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    proj_data = [
        ("Mois 1", "2 000", "60", "95 €", "1 140 €"),
        ("Mois 6", "15 000", "450", "712 €", "8 550 €"),
        ("Mois 12", "30 000", "900", "1 425 €", "17 100 €"),
        ("Mois 18", "48 000", "1 350", "2 137 €", "25 650 €"),
        ("Mois 24", "70 000", "2 100", "3 325 €", "39 900 €"),
    ]

    for i, row_data in enumerate(proj_data):
        row = table10.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            if j == 4:  # ARR
                row.cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Hypothèses
    hyp = doc.add_paragraph()
    run = hyp.add_run("Hypothèses : ")
    run.bold = True
    hyp.add_run("2 000 downloads/mois, taux de conversion 3%, churn annuel 25%")
    hyp.paragraph_format.space_after = Pt(6)

    # Break-even
    be = doc.add_paragraph()
    run = be.add_run("Break-even : ")
    run.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    be.add_run("181 abonnés actifs (coûts infrastructure ~286€/mois)")

    doc.add_page_break()

    # Coûts
    costs_title = doc.add_paragraph()
    run = costs_title.add_run("Estimation des coûts mensuels")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table11 = doc.add_table(rows=7, cols=2)
    table11.style = 'Table Grid'

    header_row11 = table11.rows[0]
    header_row11.cells[0].text = "Poste"
    header_row11.cells[1].text = "Coût mensuel"
    for cell in header_row11.cells:
        set_cell_shading(cell, '2563EB')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    costs = [
        ("Infrastructure Azure", "256 €"),
        ("Apple Developer Program", "8 € (99€/an)"),
        ("Google Play Console", "2 € (25€ amorti)"),
        ("Domaine & services", "20 €"),
        ("TOTAL", "286 €"),
    ]

    for i, (poste, cout) in enumerate(costs):
        row = table11.rows[i + 1]
        row.cells[0].text = poste
        row.cells[1].text = cout
        if poste == "TOTAL":
            set_cell_shading(row.cells[0], '2563EB')
            set_cell_shading(row.cells[1], '2563EB')
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True

    doc.add_paragraph()

    # Note scalabilité
    scale = doc.add_paragraph()
    run = scale.add_run("💡 Scalabilité : ")
    run.bold = True
    scale.add_run("Les coûts Azure sont serverless et s'adaptent automatiquement à la charge. À 100 000 utilisateurs, estimation ~900€/mois.")

    doc.add_page_break()

    # ==========================================
    # 6. PLANNING & INVESTISSEMENT
    # ==========================================

    section_title = doc.add_paragraph()
    run = section_title.add_run("6. PLANNING & INVESTISSEMENT")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(37, 99, 235)
    add_horizontal_line(section_title)

    doc.add_paragraph()

    # Phases
    phases_title = doc.add_paragraph()
    run = phases_title.add_run("Phases de développement")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table12 = doc.add_table(rows=7, cols=3)
    table12.style = 'Table Grid'

    header_row12 = table12.rows[0]
    for i, h in enumerate(["Phase", "Objectif", "Livrables clés"]):
        header_row12.cells[i].text = h
        set_cell_shading(header_row12.cells[i], '2563EB')
        for run in header_row12.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    phases = [
        ("1. Foundation", "MVP Core", "Architecture, NFC basique, Auth"),
        ("2. Core Features", "NFC Complet", "Lecture/écriture avancée, copie"),
        ("3. Business Cards", "Cartes de visite", "Création, partage, QR"),
        ("4. Premium", "Monétisation", "HCE, Wallet, abonnements"),
        ("5. Polish", "Production", "Tests, i18n, stores"),
        ("6. Post-Launch", "Amélioration", "Monitoring, feedback, v1.1"),
    ]

    for i, (phase, obj, livrables) in enumerate(phases):
        row = table12.rows[i + 1]
        row.cells[0].text = phase
        row.cells[1].text = obj
        row.cells[2].text = livrables
        set_cell_shading(row.cells[0], 'F3F4F6')
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # KPIs
    kpi_title = doc.add_paragraph()
    run = kpi_title.add_run("Indicateurs de succès")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table13 = doc.add_table(rows=6, cols=3)
    table13.style = 'Table Grid'

    header_row13 = table13.rows[0]
    for i, h in enumerate(["KPI", "Objectif", "Mesure"]):
        header_row13.cells[i].text = h
        set_cell_shading(header_row13.cells[i], '10B981')
        for run in header_row13.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    kpis = [
        ("Crash-free rate", "> 99.5%", "Firebase Crashlytics"),
        ("Temps démarrage", "< 2 secondes", "App Insights"),
        ("Temps réponse API", "< 200ms (p95)", "Azure Monitor"),
        ("Taux succès scan", "> 95%", "Analytics custom"),
        ("Note stores", "> 4.5 / 5", "App Store / Play Store"),
    ]

    for i, (kpi, obj, mesure) in enumerate(kpis):
        row = table13.rows[i + 1]
        row.cells[0].text = kpi
        row.cells[1].text = obj
        row.cells[2].text = mesure
        row.cells[1].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # Risques
    risk_title = doc.add_paragraph()
    run = risk_title.add_run("Matrice des risques")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    table14 = doc.add_table(rows=6, cols=4)
    table14.style = 'Table Grid'

    header_row14 = table14.rows[0]
    for i, h in enumerate(["Risque", "Impact", "Probabilité", "Mitigation"]):
        header_row14.cells[i].text = h
        set_cell_shading(header_row14.cells[i], 'DC2626')
        for run in header_row14.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    risks = [
        ("Limitations iOS NFC", "Élevé", "Haute", "Alternatives + communication claire"),
        ("Rejet App Store", "Élevé", "Moyenne", "Respect strict guidelines"),
        ("Failles sécurité", "Élevé", "Faible", "Audit + pentesting"),
        ("Dépassement budget Azure", "Moyen", "Moyenne", "Alertes + auto-scaling limité"),
        ("Faible adoption", "Élevé", "Moyenne", "ASO + version gratuite attractive"),
    ]

    for i, (risque, impact, proba, mitigation) in enumerate(risks):
        row = table14.rows[i + 1]
        row.cells[0].text = risque
        row.cells[1].text = impact
        row.cells[2].text = proba
        row.cells[3].text = mitigation

        # Coloration selon impact
        if impact == "Élevé":
            set_cell_shading(row.cells[1], 'FEE2E2')
        elif impact == "Moyen":
            set_cell_shading(row.cells[1], 'FEF3C7')

    doc.add_page_break()

    # ==========================================
    # 7. POURQUOI NFC PRO ?
    # ==========================================

    section_title = doc.add_paragraph()
    run = section_title.add_run("7. POURQUOI NFC PRO ?")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(37, 99, 235)
    add_horizontal_line(section_title)

    doc.add_paragraph()

    # Avantages
    adv_title = doc.add_paragraph()
    run = adv_title.add_run("Nos avantages compétitifs")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()

    advantages = [
        ("🎯 Exhaustivité", "Support de 30+ types de puces, y compris les plus récentes (NTAG 424 DNA, DESFire EV3)"),
        ("🔄 Tout-en-un", "Une seule app pour lire, écrire, copier, émuler et gérer ses cartes de visite"),
        ("💳 Intégration Wallet", "Seule app du marché à proposer l'export natif vers Google et Apple Wallet"),
        ("☁️ Cloud sync", "Synchronisation transparente entre tous vos appareils"),
        ("🔐 Sécurité Pro", "Chiffrement de bout en bout, conformité RGPD, audit sécurité"),
        ("🎨 UX Premium", "Interface intuitive pensée pour les professionnels exigeants"),
    ]

    for emoji_title, desc in advantages:
        p = doc.add_paragraph()
        run = p.add_run(emoji_title)
        run.bold = True
        run.font.size = Pt(12)

        p2 = doc.add_paragraph()
        p2.add_run(desc)
        p2.paragraph_format.left_indent = Cm(1)
        p2.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    # Call to action
    cta_title = doc.add_paragraph()
    cta_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cta_title.add_run("Prêt à révolutionner votre utilisation du NFC ?")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(37, 99, 235)

    doc.add_paragraph()

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run("Contactez-nous pour démarrer le projet")
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("─" * 40)
    run.font.color.rgb = RGBColor(209, 213, 219)

    doc.add_paragraph()

    thanks = doc.add_paragraph()
    thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = thanks.add_run("Merci de votre attention")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()

    copyright_p = doc.add_paragraph()
    copyright_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = copyright_p.add_run("© 2025 NFC Pro — Document confidentiel")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(156, 163, 175)

    # Sauvegarde
    doc.save('C:/Users/jcpas/NFCPro/NFC_Pro_Presentation.docx')
    print("Document cree avec succes : NFC_Pro_Presentation.docx")

if __name__ == "__main__":
    create_document()
